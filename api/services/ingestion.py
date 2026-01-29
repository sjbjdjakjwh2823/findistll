"""
FinDistill File Ingestion Service

Handles parsing of various file formats using Gemini API via HTTP.

[Input Formats - Optimized Processing]
- PDF (.pdf): Gemini multimodal for complex table extraction
- Images (.jpg, .png, .tiff, .webp, .heic): Gemini OCR + structure understanding
- Word (.docx): python-docx for text extraction + Gemini for summarization
- HWP (.hwpx): XML extraction + Gemini for summarization  
- Excel (.xlsx, .xls): openpyxl for structured data + Gemini for summary
- CSV (.csv): Python csv module + Gemini for summary

[Output Formats]
- JSONL: For LLM fine-tuning (instruction-response pairs)
- Markdown: For RAG systems (preserves document hierarchy)
"""

import io
import json
import csv
import base64
import zipfile
import xml.etree.ElementTree as ET
from typing import Dict, Any, List, Optional
import os
import logging
import copy
import httpx

# Standard imports assuming requirements.txt is satisfied
import docx

# Configure logger
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Gemini API configuration
GEMINI_API_BASE = "https://generativelanguage.googleapis.com/v1beta"


class GeminiClient:
    """Simple Gemini API client using httpx."""
    
    def __init__(self):
        self.api_key = os.getenv("GEMINI_API_KEY", "")
        self.model = "gemini-2.0-flash"
        
    async def generate_content(self, contents: list, response_mime_type: str = None) -> str:
        """Generate content using Gemini API."""
        if not self.api_key:
            logger.warning("Gemini API Key missing. Returning mock response.")
            if response_mime_type == "application/json":
                return json.dumps({
                    "title": "Mock Title",
                    "summary": "Mock Summary (No API Key)",
                    "tables": [],
                    "key_metrics": {},
                    "currency": "USD",
                    "date_range": "2023-2024"
                })
            return "Mock Content (No API Key)"

        url = f"{GEMINI_API_BASE}/models/{self.model}:generateContent?key={self.api_key}"
        
        generation_config = {}
        if response_mime_type:
            generation_config["responseMimeType"] = response_mime_type
        
        payload = {"contents": contents}
        if generation_config:
            payload["generationConfig"] = generation_config
        
        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.post(url, json=payload)
            response.raise_for_status()
            data = response.json()
            
        # Extract text from response
        candidates = data.get("candidates", [])
        if candidates:
            parts = candidates[0].get("content", {}).get("parts", [])
            if parts:
                return parts[0].get("text", "")
        return ""
    
    async def generate_with_file(self, file_content: bytes, mime_type: str, prompt: str, response_mime_type: str = None) -> str:
        """Generate content with inline file data."""
        if not self.api_key:
            logger.warning("Gemini API Key missing. Returning mock response.")
            return json.dumps({
                "title": "Mock Analysis",
                "summary": "Mock Summary (No API Key)",
                "tables": [],
                "key_metrics": {},
                "currency": "USD",
                "date_range": "2023-2024"
            })

        # Encode file as base64
        file_b64 = base64.b64encode(file_content).decode("utf-8")
        
        contents = [{
            "parts": [
                {
                    "inline_data": {
                        "mime_type": mime_type,
                        "data": file_b64
                    }
                },
                {"text": prompt}
            ]
        }]
        
        return await self.generate_content(contents, response_mime_type)


class FileIngestionService:
    """Service for ingesting and parsing various file formats."""
    
    # Comprehensive file format support
    SUPPORTED_FORMATS = {
        # PDF - Gemini multimodal for complex layout
        'application/pdf': 'pdf',
        
        # Images - Gemini OCR and structure recognition
        'image/png': 'image',
        'image/jpeg': 'image',
        'image/tiff': 'image',
        'image/webp': 'image',
        'image/heic': 'image',
        
        # Word documents - python-docx extraction
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        
        # HWP (Korean word processor) - XML extraction from HWPX
        'application/hwp+zip': 'hwpx',
        'application/x-hwpx': 'hwpx',
        
        # Excel - openpyxl for structured data
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'excel',
        'application/vnd.ms-excel': 'excel',
        
        # CSV - Python csv module
        'text/csv': 'csv',
        
        # XML/XBRL - Enterprise financial data
        'application/xml': 'xbrl',
        'text/xml': 'xbrl',
        'application/xbrl+xml': 'xbrl',
    }

    def __init__(self):
        self.gemini = GeminiClient()

    async def process_file(
        self, 
        file_content: bytes, 
        filename: str, 
        mime_type: str
    ) -> Dict[str, Any]:
        """Process a file and extract structured financial data."""
        file_type = self.SUPPORTED_FORMATS.get(mime_type, 'unknown')
        
        # Auto-detect XBRL by filename extension
        if filename.lower().endswith(('.xbrl', '.xml')) and file_type == 'unknown':
            file_type = 'xbrl'
        
        if file_type == 'csv':
            return await self._process_csv(file_content, filename)
        elif file_type == 'excel':
            return await self._process_excel(file_content, filename)
        elif file_type == 'xbrl':
            return await self._process_xbrl(file_content, filename)
        elif file_type == 'docx':
            return await self._process_docx(file_content, filename)
        elif file_type == 'hwpx':
            return await self._process_hwpx(file_content, filename)
        elif file_type in ('pdf', 'image'):
            return await self._process_with_gemini(file_content, filename, mime_type)
        else:
            raise ValueError(f"Unsupported file type: {mime_type}")

    async def _process_docx(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process Word document using python-docx and Gemini."""
        # Using global import docx
        doc = docx.Document(io.BytesIO(content))
        full_text = []
        
        # Extract parsing structure
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
                
        # Basic table extraction (heuristic)
        tables_data = []
        for table in doc.tables:
            t_rows = []
            for row in table.rows:
                t_rows.append([cell.text.strip() for cell in row.cells])
            if t_rows:
                tables_data.append({
                    "name": f"Table_{len(tables_data)+1}",
                    "headers": t_rows[0],
                    "rows": t_rows[1:]
                })
        
        text_content = "\n".join(full_text)
        
        # Use Gemini to structure and analyze the extracted text
        return await self._analyze_text_with_gemini(text_content, filename, "docx", tables_data)

    async def _process_hwpx(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process HWPX file by extracting text from XML."""
        text_content = ""
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as zf:
                # HWPX structure usually has content in Contents/section0.xml
                # But we should search for all section XMLs
                section_files = [f for f in zf.namelist() if f.startswith('Contents/section') and f.endswith('.xml')]
                
                for section_file in sorted(section_files):
                    xml_data = zf.read(section_file)
                    root = ET.fromstring(xml_data)
                    
                    # Extract text from <hp:t> tags (HWPX text tag)
                    # Note: Namespace handling might be needed, but simple search often works
                    # Let's try searching for all text nodes
                    texts = [elem.text for elem in root.iter() if elem.text and elem.text.strip()]
                    text_content += "\n".join(texts) + "\n\n"
                    
        except Exception as e:
            logger.error(f"Error parsing HWPX: {e}")
            raise ValueError(f"Failed to parse HWPX file: {e}")

        if not text_content:
            text_content = "No extractable text found in HWPX file."
            
        return await self._analyze_text_with_gemini(text_content, filename, "hwpx")

    async def _analyze_text_with_gemini(self, text: str, filename: str, file_type: str, pre_extracted_tables: List[Dict] = None) -> Dict[str, Any]:
        """Helper to send text content to Gemini for structuring."""
        prompt = f'''
        Analyze the following text extracted from a {file_type} document ({filename}).
        Extract structured financial data into JSON.
        
        Text Content:
        {text[:30000]}  # Limit context window if necessary, though Gemini Flash has 1M window
        
        Requirements:
        1. Identify the document title and provide a detailed summary.
        2. Identify key financial metrics (revenue, profit, ratios).
        3. If there are tables explicitly mentioned or structured in text, reconstruct them.
        4. Output strictly in JSON.
        5. IMPORTANT: All output must be in English. Translate if necessary.
        
        Output JSON format:
        {{
            "title": "Document Title",
            "summary": "Detailed summary",
            "tables": [
                {{
                    "name": "Table Name",
                    "headers": ["Col1", "Col2"],
                    "rows": [["Val1", "Val2"]]
                }}
            ],
            "key_metrics": {{ "metric": "value" }},
            "currency": "KRW/USD",
            "date_range": "YYYY-MM-DD"
        }}
        '''
        
        response_text = await self.gemini.generate_content(
            [{"parts": [{"text": prompt}]}], 
            "application/json"
        )
        
        try:
            result = json.loads(response_text)
        except json.JSONDecodeError:
            # Fallback if valid JSON isn't returned
            result = {
                "title": filename, 
                "summary": "AI processing error", 
                "tables": [], 
                "key_metrics": {}
            }
            
        # Merge pre-extracted tables if AI didn't find them but we did (e.g. from DOCX)
        if pre_extracted_tables and not result.get("tables"):
            result["tables"] = pre_extracted_tables
            
        result["metadata"] = {
            "file_type": file_type,
            "processed_by": "gemini-2.0-flash-text"
        }
        
        # Ensure reasoning_qa exists
        if "reasoning_qa" not in result:
            result["reasoning_qa"] = [
                {
                    "question": "Provide a summary of the document.",
                    "response": result.get("summary", "No summary available."),
                    "type": "summary"
                }
            ]
            
            # Add table analysis
            for table in result.get("tables", []):
                t_name = table.get("name", "Unknown Table")
                result["reasoning_qa"].append({
                    "question": f"Analyze the table '{t_name}'.",
                    "response": f"Table '{t_name}' contains {len(table.get('rows', []))} rows of data.",
                    "type": "table_analysis"
                })
        
        return result

    async def _process_xbrl(self, content: bytes, filename: str) -> Dict[str, Any]:
        """
        Process XBRL/XML file using STRUCTURAL Tree Parser.
        
        [Workflow]
        1. Parse structure with XML Tree Parser
        2. Standardize units via ScaleProcessor ($B)
        3. Filter Contexts (CY/PY mapping)
        4. Infer reasoning Q&A in CoT format
        
        Returns: XBRLIntelligenceResult converted to standard format
        """
        from .xbrl_semantic_engine import XBRLSemanticEngine
        
        # Automatic detection of label linkbase (filename-based)
        label_content = None
        base_name = filename.replace('.htm.xml', '').replace('.xml', '').replace('.xbrl', '')
        
        # NOTE: In a real implementation, _lab.xml files should also be uploaded.
        # Currently processing with instance file only (using CoreFinancialConcepts fallback).
        
        # Initialize XBRLSemanticEngine and perform structural parsing.
        engine = XBRLSemanticEngine(
            company_name="",  # Extracted during parsing
            fiscal_year="",    # Extracted during parsing
            file_path=filename
        )
        
        try:
            result = engine.process_joint(
                label_content=label_content,
                instance_content=content
            )
            
            if not result.success:
                return {
                    "title": f"XBRL Parsing Failed: {filename}",
                    "summary": result.parse_summary,
                    "tables": [],
                    "key_metrics": {},
                    "facts": [],
                    "parse_log": result.errors,
                    "metadata": {
                        "file_type": "xbrl",
                        "processed_by": "xbrl-semantic-engine-v11.5",
                        "error": True
                    }
                }
            
            # Convert SemanticFacts to standard format
            facts_list = []
            for fact in result.facts:
                facts_list.append({
                    "concept": fact.concept,
                    "label": fact.label,
                    "value": str(fact.value),
                    "raw_value": fact.raw_value,
                    "unit": fact.unit,
                    "period": fact.period,
                    "is_consolidated": fact.is_consolidated,
                    "decimals": fact.decimals
                })
            
            # Convert to table format
            tables = self._build_financial_tables(facts_list)

            # CRITICAL DEBUG: Verify Data Pipe before return
            final_qa = copy.deepcopy(result.reasoning_qa)
            
            # [Strict Handover] Defensive Check
            if not isinstance(final_qa, list):
                logger.error(f"CRITICAL HANDOVER ERROR: reasoning_qa is not a list! Type: {type(final_qa)}")
                final_qa = []  # Emergency reset
            
            qa_count = len(final_qa)
            logger.info(f"V11.5 DATA RECOVERED: [{qa_count}] ROWS READY")
            logger.info(f"TRACE 5: Final data count being sent to Exporter: {qa_count}")
            
            return {
                "title": f"XBRL: {result.company_name or filename}",
                "summary": result.parse_summary,
                "tables": tables,
                "key_metrics": result.key_metrics,
                "facts": facts_list,
                "reasoning_qa": final_qa, # CRITICAL: Deep copy pass-through
                "jsonl_data": result.jsonl_data,  # v11.5: Pass pre-verified JSONL
                "financial_report_md": result.financial_report_md,
                "parse_log": [],
                "metadata": {
                    "file_type": "xbrl",
                    "company": result.company_name,
                    "fiscal_year": result.fiscal_year,
                    "fact_count": len(facts_list),
                    "processed_by": "xbrl-semantic-engine-v11.5"
                }
            }
            
        except Exception as e:
            logger.error(f"XBRL Parsing totally failed: {e}")
            return {
                "title": f"XBRL Parsing Error: {filename}",
                "summary": f"Fatal error: {str(e)}",
                "tables": [],
                "key_metrics": {},
                "facts": [],
                "metadata": {
                    "file_type": "xbrl",
                    "error": True,
                    "processed_by": "none"
                }
            }
    
    def _build_financial_tables(self, facts: List[Dict]) -> List[Dict]:
        """Convert facts into financial table format."""
        
        # Balance Sheet items
        balance_sheet = {
            "name": "Statement of Financial Position (Balance Sheet)",
            "headers": ["Account", "Amount ($B)", "Period"],
            "rows": []
        }
        
        # Income Statement items
        income_statement = {
            "name": "Statement of Comprehensive Income (Income Statement)",
            "headers": ["Account", "Amount ($B)", "Period"],
            "rows": []
        }
        
        for fact in facts:
            label = fact.get("label", "")
            value = fact.get("value", "")
            period = fact.get("period", "")
            
            row = [label, value, period]
            
            # Simple categorization based on core keywords
            if any(k in label.lower() for k in ['asset', 'liabilit', 'equity']):
                balance_sheet["rows"].append(row)
            elif any(k in label.lower() for k in ['revenue', 'profit', 'income', 'loss', 'expens']):
                income_statement["rows"].append(row)
        
        tables = []
        if balance_sheet["rows"]:
            tables.append(balance_sheet)
        if income_statement["rows"]:
            tables.append(income_statement)
            
        return tables

    async def _process_csv(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process CSV file using standard csv module."""
        text_content = None
        
        # Try different encodings
        for encoding in ['utf-8', 'cp949', 'euc-kr', 'latin1']:
            try:
                text_content = content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        
        if text_content is None:
            raise ValueError("Could not decode CSV file with any supported encoding")
        
        # Parse CSV
        f = io.StringIO(text_content)
        reader = csv.reader(f)
        rows = list(reader)
        
        if not rows:
            return {"title": filename, "summary": "Empty CSV file", "tables": []}
            
        headers = rows[0]
        data_rows = rows[1:]
        
        tables = [{
            "name": filename,
            "headers": headers,
            "rows": data_rows
        }]
        
        # Generate summary using Gemini
        sample_lines = rows[:10]
        sample_data = "\n".join([",".join(map(str, row)) for row in sample_lines])
        summary = await self._generate_summary(sample_data)
        
        # Generate Reasoning QA for JSONL
        reasoning_qa = [
            {
                "question": "Provide a summary of the provided CSV data.",
                "response": summary,
                "type": "summary"
            }
        ]
        
        # Add metric analysis to reasoning_qa
        metrics = self._extract_metrics(headers, data_rows)
        for k, v in metrics.items():
            reasoning_qa.append({
                "question": f"Analyze the metric '{k}' from the dataset.",
                "response": f"The calculated value for {k} is {v}. This metric provides insight into the dataset's aggregate performance.",
                "type": "metric_analysis"
            })
        
        return {
            "title": filename,
            "summary": summary,
            "tables": tables,
            "key_metrics": metrics,
            "reasoning_qa": reasoning_qa,
            "metadata": {
                "file_type": "csv",
                "row_count": len(rows),
                "column_count": len(headers)
            }
        }
    
    async def _process_excel(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process Excel file using openpyxl."""
        from openpyxl import load_workbook
        
        wb = load_workbook(filename=io.BytesIO(content), data_only=True)
        tables = []
        all_metrics = {}
        
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = list(ws.values)
            
            if not rows:
                continue
                
            headers = list(rows[0]) if rows else []
            data_rows = [list(row) for row in rows[1:]] if len(rows) > 1 else []
            
            # Simple conversion of None to empty string
            headers = [str(h) if h is not None else "" for h in headers]
            data_rows = [[cell if cell is not None else "" for cell in row] for row in data_rows]
            
            tables.append({
                "name": sheet_name,
                "headers": headers,
                "rows": data_rows
            })
            
            # Merge metrics
            sheet_metrics = self._extract_metrics(headers, data_rows)
            all_metrics.update({f"{sheet_name}_{k}": v for k, v in sheet_metrics.items()})
        
        # Generate summary
        sample_text = "\n".join([f"{t['name']}: {t['headers']}" for t in tables[:3]])
        summary = await self._generate_summary(sample_text)
        
        # Generate Reasoning QA for JSONL
        reasoning_qa = [
            {
                "question": "Provide a summary of the Excel data.",
                "response": summary,
                "type": "summary"
            }
        ]
        
        # Add metric analysis
        for k, v in all_metrics.items():
            reasoning_qa.append({
                "question": f"Analyze the metric '{k}' from the workbook.",
                "response": f"The calculated value for {k} is {v}.",
                "type": "metric_analysis"
            })
        
        return {
            "title": filename,
            "summary": summary,
            "tables": tables,
            "key_metrics": all_metrics,
            "reasoning_qa": reasoning_qa,
            "metadata": {
                "file_type": "excel",
                "sheet_count": len(tables)
            }
        }
    
    async def _process_with_gemini(self, content: bytes, filename: str, mime_type: str) -> Dict[str, Any]:
        """Process PDF/Image using Gemini multimodal via HTTP."""
        prompt = '''
        Analyze this financial document thoroughly. Extract ALL data into structured JSON.
        
        Requirements:
        1. Identify the document title and provide a detailed summary
        2. Extract ALL tables with proper headers and data
        3. Identify key financial metrics (revenue, profit, growth rates, ratios)
        4. Note any currency units (KRW, USD, etc.)
        5. Identify date references and time periods
        6. IMPORTANT: All output must be in English. Translate if necessary.
        
        Output JSON format:
        {
            "title": "Document Title",
            "summary": "Detailed summary of the document content",
            "tables": [
                {
                    "name": "Table Name",
                    "headers": ["Column1", "Column2", ...],
                    "rows": [["Value1", "Value2", ...], ...]
                }
            ],
            "key_metrics": {
                "metric_name": "value with unit"
            },
            "currency": "KRW or USD",
            "date_range": "YYYY-MM-DD to YYYY-MM-DD"
        }
        '''
        
        response_text = await self.gemini.generate_with_file(
            content, mime_type, prompt, "application/json"
        )
        
        result = json.loads(response_text)
        result["metadata"] = {
            "file_type": "pdf" if "pdf" in mime_type else "image",
            "processed_by": "gemini-2.0-flash"
        }
        
        # Ensure reasoning_qa exists
        if "reasoning_qa" not in result:
            result["reasoning_qa"] = [
                {
                    "question": "Provide a summary of the document.",
                    "response": result.get("summary", "No summary available."),
                    "type": "summary"
                }
            ]
            
            for k, v in result.get("key_metrics", {}).items():
                result["reasoning_qa"].append({
                    "question": f"What is the value of {k}?",
                    "response": f"The value of {k} is {v}.",
                    "type": "metric_fact"
                })
        
        return result
    
    async def _generate_summary(self, data_sample: str) -> str:
        """Generate a summary using Gemini."""
        prompt = f"Summarize the key findings of the following data in 2-3 sentences:\n\n{data_sample}"
        
        contents = [{"parts": [{"text": prompt}]}]
        return await self.gemini.generate_content(contents)
    
    def _extract_metrics(self, headers: List[str], rows: List[List[Any]]) -> Dict[str, Any]:
        """Extract simple metrics from headers and rows without pandas."""
        metrics = {}
        if not headers or not rows:
            return metrics

        # Find potential numeric columns (simple heuristic)
        numeric_indices = []
        for i in range(len(headers)):
            # Check first few non-empty rows
            is_numeric = False
            for row in rows[:5]:
                if i < len(row) and isinstance(row[i], (int, float)):
                   is_numeric = True
                   break
                elif i < len(row) and isinstance(row[i], str) and row[i].replace('.','',1).isdigit():
                   is_numeric = True
                   break
            if is_numeric:
                numeric_indices.append(i)
        
        # Calculate sums/averages for first few numeric columns
        for i in numeric_indices[:5]:
            col_name = str(headers[i])
            values = []
            for row in rows:
                if i < len(row):
                    val = row[i]
                    try:
                        if isinstance(val, (int, float)):
                            values.append(float(val))
                        elif isinstance(val, str):
                            # clean string
                            clean_val = val.replace(',', '').replace(' ', '')
                            if clean_val:
                                values.append(float(clean_val))
                    except ValueError:
                        continue
            
            if values:
                metrics[f"{col_name}_total"] = sum(values)
                metrics[f"{col_name}_avg"] = sum(values) / len(values)
                
        return metrics

ingestion_service = FileIngestionService()
